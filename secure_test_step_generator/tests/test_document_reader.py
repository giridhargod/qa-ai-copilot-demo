import io
import os
import tempfile

import docx
import pytest
from PIL import Image

from ..document_reader import extract_evidence

FIXTURES_DIR = os.path.join(os.path.dirname(__file__), "fixtures")


def _png_bytes(size=(50, 50), color="white"):
    buf = io.BytesIO()
    Image.new("RGB", size, color=color).save(buf, format="PNG")
    return buf.getvalue()


def _jpeg_bytes(size=(50, 50), color="red"):
    buf = io.BytesIO()
    Image.new("RGB", size, color=color).save(buf, format="JPEG")
    return buf.getvalue()


def test_extracts_docx_screenshots_in_order_with_notes():
    items = extract_evidence(os.path.join(FIXTURES_DIR, "sample.docx"))

    assert len(items) == 2
    assert items[0].order == 0
    assert items[1].order == 1

    assert "claim intake screen" in items[0].notes_text.lower()
    assert "CL-99231" in items[0].notes_text

    assert "Click Submit" in items[1].notes_text

    assert items[0].image_bytes.startswith(b"\x89PNG")
    assert items[1].image_bytes.startswith(b"\x89PNG")


def test_extracts_pdf_screenshot_with_page_notes():
    items = extract_evidence(os.path.join(FIXTURES_DIR, "sample.pdf"))

    assert len(items) == 1
    assert "employer verification page" in items[0].notes_text.lower()
    assert "EMP-4471" in items[0].notes_text
    assert items[0].source_ref == "page 1"
    assert len(items[0].image_bytes) > 0


def test_unsupported_extension_raises():
    with pytest.raises(ValueError):
        extract_evidence("notes.txt")


def test_extracts_screenshot_pasted_inside_a_table_cell():
    # Regression: document.paragraphs (the "obvious" python-docx API)
    # only returns body-level paragraphs -- a screenshot pasted inside
    # a table cell (a common QA test-case template layout) used to be
    # silently dropped with zero warning.
    items = extract_evidence(os.path.join(FIXTURES_DIR, "sample_with_table.docx"))

    assert len(items) == 1
    assert "Intro note before the table" in items[0].notes_text
    assert "Step 1: Open the claim record" in items[0].notes_text
    assert "table" in items[0].source_ref
    assert items[0].image_bytes.startswith(b"\x89PNG")


def test_missing_docx_raises_file_not_found_not_a_corruption_error():
    # Regression: a typo'd/missing .docx path used to be caught inside
    # _extract_from_docx and reported as a misleading "may be corrupted"
    # ValueError -- a missing .pdf already raised FileNotFoundError
    # naturally (via pypdf's own file open), so the two extensions gave
    # inconsistent errors for the same root cause.
    with pytest.raises(FileNotFoundError):
        extract_evidence(os.path.join(FIXTURES_DIR, "does_not_exist.docx"))


def test_missing_pdf_raises_file_not_found():
    with pytest.raises(FileNotFoundError):
        extract_evidence(os.path.join(FIXTURES_DIR, "does_not_exist.pdf"))


def test_corrupt_docx_raises_clean_value_error_not_library_exception():
    with pytest.raises(ValueError, match="Could not read"):
        extract_evidence(os.path.join(FIXTURES_DIR, "corrupt.docx"))


def test_corrupt_pdf_raises_clean_value_error_not_library_exception():
    with pytest.raises(ValueError, match="Could not read"):
        extract_evidence(os.path.join(FIXTURES_DIR, "corrupt.pdf"))


def test_encrypted_pdf_raises_clean_value_error_not_library_exception():
    with pytest.raises(ValueError, match="password-protected"):
        extract_evidence(os.path.join(FIXTURES_DIR, "encrypted.pdf"))


def test_docx_with_no_screenshots_returns_empty_list_not_an_error():
    document = docx.Document()
    document.add_paragraph("Just some notes, no screenshots pasted at all.")
    with tempfile.TemporaryDirectory() as tmp_dir:
        path = os.path.join(tmp_dir, "no_images.docx")
        document.save(path)
        assert extract_evidence(path) == []


def test_multiple_images_in_a_single_paragraph_are_both_captured():
    document = docx.Document()
    paragraph = document.add_paragraph()
    run1 = paragraph.add_run()
    run1.add_picture(io.BytesIO(_jpeg_bytes()), width=docx.shared.Inches(0.5))
    run2 = paragraph.add_run()
    run2.add_picture(io.BytesIO(_png_bytes()), width=docx.shared.Inches(0.5))

    with tempfile.TemporaryDirectory() as tmp_dir:
        path = os.path.join(tmp_dir, "multi_image.docx")
        document.save(path)
        items = extract_evidence(path)

        assert len(items) == 2
        assert items[0].image_bytes.startswith(b"\xff\xd8\xff")  # JPEG magic bytes
        assert items[1].image_bytes.startswith(b"\x89PNG")
