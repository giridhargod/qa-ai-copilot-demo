"""Document stage: read a Word/PDF file, extract screenshots with the
tester's surrounding notes, in original document order.

Deliberately does not import qa-ai-copilot's services/file_service.py
-- this package stays self-contained so it can be cloned/copied on
its own (Architect decision #1). The extraction logic below is a
focused subset of what file_service.py does (it only extracts text),
extended to also pull the embedded screenshots themselves.
"""
import os

import docx
from docx.opc.exceptions import PackageNotFoundError
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader
from pypdf.errors import FileNotDecryptedError, PyPdfError

from .models import EvidenceItem

SUPPORTED_EXTENSIONS = (".docx", ".pdf")


def _iter_block_items(parent):
    """Yield Paragraph/Table children of `parent` in true document
    order. `parent` is either the Document itself or a table cell.

    document.paragraphs / document.tables (the "obvious" python-docx
    API) only return body-level items -- a screenshot pasted inside a
    table cell (a common QA test-case template layout: "Step | Screenshot
    | Notes" columns) is invisible to them and gets silently dropped.
    This walks the underlying XML body directly instead, recursing into
    table cells so nothing is missed.
    """
    parent_elm = parent.element.body if isinstance(parent, docx.document.Document) else parent._tc

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def extract_evidence(file_path: str) -> list[EvidenceItem]:
    """Dispatch by extension. Returns EvidenceItem[] in document order."""
    ext = os.path.splitext(file_path)[1].lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported input document type '{ext}'. "
            f"Supported types: {', '.join(SUPPORTED_EXTENSIONS)}"
        )

    if not os.path.isfile(file_path):
        # python-docx raises its own PackageNotFoundError for a missing
        # file, which _extract_from_docx below would otherwise turn into
        # a misleading "may be corrupted" ValueError -- check up front so
        # a missing .docx and a missing .pdf fail the same clear way
        # (pypdf already raises FileNotFoundError natively for its path).
        raise FileNotFoundError(f"Input document not found: {file_path}")

    if ext == ".docx":
        return _extract_from_docx(file_path)
    return _extract_from_pdf(file_path)


def _extract_from_docx(file_path: str) -> list[EvidenceItem]:
    try:
        document = docx.Document(file_path)
    except PackageNotFoundError as exc:
        raise ValueError(
            f"Could not read '{file_path}' as a Word document -- it may "
            f"be corrupted, password-protected, or not actually a .docx "
            f"file ({exc})."
        ) from exc

    items: list[EvidenceItem] = []
    notes_buffer: list[str] = []
    order_counter = [0]

    def handle_paragraph(paragraph: Paragraph, source_ref: str) -> None:
        image_rids = [
            blip.get(qn("r:embed"))
            for run in paragraph.runs
            for blip in run._element.findall(".//" + qn("a:blip"))
            if blip.get(qn("r:embed"))
        ]

        if image_rids:
            for rid in image_rids:
                image_part = paragraph.part.related_parts[rid]
                items.append(
                    EvidenceItem(
                        order=order_counter[0],
                        image_bytes=image_part.blob,
                        notes_text="\n".join(notes_buffer).strip(),
                        source_ref=source_ref,
                    )
                )
                order_counter[0] += 1
            notes_buffer.clear()
        else:
            text = paragraph.text.strip()
            if text:
                notes_buffer.append(text)

    def walk(parent, ref_label: str) -> None:
        for block_index, block in enumerate(_iter_block_items(parent)):
            if isinstance(block, Paragraph):
                handle_paragraph(block, f"{ref_label} {block_index}")
            elif isinstance(block, Table):
                for row_index, row in enumerate(block.rows):
                    for cell_index, cell in enumerate(row.cells):
                        walk(cell, f"{ref_label} table row {row_index} cell {cell_index}")

    walk(document, "paragraph")
    return items


def _extract_from_pdf(file_path: str) -> list[EvidenceItem]:
    """pypdf (pure Python, no native/compiled dependency) doesn't expose
    per-image bounding boxes, so unlike the docx path this can only
    interleave notes at page granularity, not paragraph granularity:
    a page's extracted text becomes the notes for the first screenshot
    found on that page; any additional screenshots on the same page
    get an empty notes_text (to avoid duplicating the same note block
    across multiple rows) but keep the same page source_ref.
    """
    try:
        reader = PdfReader(file_path)
        if reader.is_encrypted:
            raise ValueError(
                f"'{file_path}' is a password-protected PDF. Remove the "
                f"password before running this tool -- encrypted PDFs "
                f"are not supported."
            )

        items: list[EvidenceItem] = []
        order = 0

        for page_index, page in enumerate(reader.pages):
            page_text = (page.extract_text() or "").strip()
            page_images = list(page.images)

            for image_index, image_file in enumerate(page_images):
                items.append(
                    EvidenceItem(
                        order=order,
                        image_bytes=image_file.data,
                        notes_text=page_text if image_index == 0 else "",
                        source_ref=f"page {page_index + 1}",
                    )
                )
                order += 1
    except FileNotDecryptedError as exc:
        raise ValueError(
            f"'{file_path}' is a password-protected PDF. Remove the "
            f"password before running this tool -- encrypted PDFs are "
            f"not supported."
        ) from exc
    except PyPdfError as exc:
        raise ValueError(
            f"Could not read '{file_path}' as a PDF -- it may be "
            f"corrupted or not actually a .pdf file ({exc})."
        ) from exc

    return items
